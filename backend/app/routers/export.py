from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from docx import Document
import os
from .. import database, models, auth

router = APIRouter(prefix="/export", tags=["Export"])

EXPORT_DIR = "exports"
os.makedirs(EXPORT_DIR, exist_ok=True)

@router.get("/{project_id}")
def export_project(project_id: int, db: Session = Depends(database.get_db), current_user: models.User = Depends(auth.get_current_user)):
    project = db.query(models.Project).filter(models.Project.id == project_id, models.Project.owner_id == current_user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    doc = Document()
    doc.add_heading(project.title, 0)
    doc.add_paragraph(f"Type: {project.project_type}\n")

    sections = db.query(models.DocumentSection).filter(models.DocumentSection.project_id == project.id).order_by(models.DocumentSection.order_index).all()
    for section in sections:
        doc.add_heading(section.title, level=1)
        doc.add_paragraph(section.content or "")

    filepath = os.path.join(EXPORT_DIR, f"{project.title.replace(' ', '_')}.docx")
    doc.save(filepath)
    return FileResponse(filepath, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=f"{project.title}.docx")
